import io

import requests
from docx import Document as DocxDocument
from langchain_core.documents import Document

from features.documents.plain_text_loader import MAX_FILE_SIZE_BYTES
from features.web_browsing.web_fetcher import DEFAULT_HEADERS
from util import log
from util.error_codes import ATTACHMENT_PROCESSING_FAILED, DOCUMENT_SEARCH_FAILED
from util.errors import ExternalServiceError


class DocxLoader:

    __job_id: str
    __document_url: str

    def __init__(self, job_id: str, document_url: str):
        self.__job_id = job_id
        self.__document_url = document_url

    def load(self) -> list[Document]:
        log.t(f"Loading DOCX document for job '{self.__job_id}'")
        try:
            raw_bytes = requests.get(self.__document_url, headers = DEFAULT_HEADERS).content
            if len(raw_bytes) > MAX_FILE_SIZE_BYTES:
                raise ExternalServiceError(
                    f"File too large for processing (>{MAX_FILE_SIZE_BYTES // (1024 * 1024)}MB)",
                    ATTACHMENT_PROCESSING_FAILED,
                )
            doc = DocxDocument(io.BytesIO(raw_bytes))
            text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
            log.t(f"Loaded DOCX document: {len(text)} characters")
            return [Document(page_content = text, metadata = {"chunk": 0})]
        except ExternalServiceError:
            raise
        except Exception as e:
            raise ExternalServiceError(
                f"Failed to load DOCX document for job '{self.__job_id}'",
                DOCUMENT_SEARCH_FAILED,
            ) from e
